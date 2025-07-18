from django.shortcuts import render
from django.http import HttpResponse
from plagiarismchecker.algorithm import main, fileSimilarity
from docx import Document
import traceback
import PyPDF2

# Home
def home(request):
    return render(request, 'pc/index.html')

# Web Search (Text Input)
def test(request):
    try:
        print("request is welcome test")
        query = request.POST.get('q', '').strip()
        if query:
            percent, link = main.findSimilarity(query)
            percent = round(percent, 2)
            print("Output: ", percent, link)
            return render(request, 'pc/index.html', {'link': link, 'percent': percent})
        else:
            return HttpResponse("No query provided", status=400)
    except Exception as e:
        print("❌ ERROR in /test/:", e)
        traceback.print_exc()
        return HttpResponse("Internal Server Error", status=500)

# Web Search (File Input)
def filetest(request):
    try:
        value = ''
        uploaded_file = request.FILES.get('docfile', None)

        if not uploaded_file:
            return HttpResponse("No file uploaded", status=400)

        filename = uploaded_file.name.lower()

        if filename.endswith(".txt"):
            value = str(uploaded_file.read(), 'utf-8', errors='ignore')

        elif filename.endswith(".docx"):
            document = Document(uploaded_file)
            for para in document.paragraphs:
                value += para.text + '\n'

        elif filename.endswith(".pdf"):
            reader = PyPDF2.PdfReader(uploaded_file)
            for page in reader.pages:
                value += page.extract_text()

        if not value.strip():
            return HttpResponse("Extracted text is empty", status=400)

        percent, link = main.findSimilarity(value)
        percent = round(percent, 2)
        print("Output:", percent, link)
        return render(request, 'pc/index.html', {'link': link, 'percent': percent})

    except Exception as e:
        print("❌ ERROR in /filetest/:", e)
        traceback.print_exc()
        return HttpResponse("Internal Server Error", status=500)

# Text Compare Page
def fileCompare(request):
    return render(request, 'pc/doc_compare.html')

# Two Text Compare (Text Input)
def twofiletest1(request):
    try:
        q1 = request.POST.get('q1', '').strip()
        q2 = request.POST.get('q2', '').strip()

        if q1 and q2:
            result = fileSimilarity.findFileSimilarity(q1, q2)
            result = round(result, 2)
            print("Output:", result)
            return render(request, 'pc/doc_compare.html', {'result': result})
        else:
            return HttpResponse("Both texts are required", status=400)

    except Exception as e:
        print("❌ ERROR in /twofiletest1/:", e)
        traceback.print_exc()
        return HttpResponse("Internal Server Error", status=500)

# Two File Compare (File Input)
def twofilecompare1(request):
    try:
        value1 = ''
        value2 = ''
        file1 = request.FILES.get('docfile1', None)
        file2 = request.FILES.get('docfile2', None)

        if not file1 or not file2:
            return HttpResponse("Both files are required", status=400)

        name1 = file1.name.lower()
        name2 = file2.name.lower()

        if name1.endswith(".txt") and name2.endswith(".txt"):
            value1 = str(file1.read(), 'utf-8', errors='ignore')
            value2 = str(file2.read(), 'utf-8', errors='ignore')

        elif name1.endswith(".docx") and name2.endswith(".docx"):
            doc1 = Document(file1)
            doc2 = Document(file2)
            value1 = '\n'.join([p.text for p in doc1.paragraphs])
            value2 = '\n'.join([p.text for p in doc2.paragraphs])

        else:
            return HttpResponse("Unsupported file types. Use .txt or .docx", status=400)

        result = fileSimilarity.findFileSimilarity(value1, value2)
        result = round(result, 2)
        print("Output:", result)
        return render(request, 'pc/doc_compare.html', {'result': result})

    except Exception as e:
        print("❌ ERROR in /twofilecompare1/:", e)
        traceback.print_exc()
        return HttpResponse("Internal Server Error", status=500)
